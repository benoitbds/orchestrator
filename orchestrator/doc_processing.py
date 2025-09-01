import io
import logging
import mimetypes
from pathlib import Path
from typing import Optional, Union

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)

class DocumentParsingError(Exception):
    """Exception raised when document parsing fails."""
    pass

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from various file types.
    
    Args:
        file_content: The raw bytes of the file
        filename: The name of the file (used for extension detection)
        
    Returns:
        Extracted text as a string
        
    Raises:
        DocumentParsingError: If the file format is unsupported or parsing fails
    """
    if not file_content:
        raise DocumentParsingError("Empty file content")
    
    # Get file extension and MIME type
    file_path = Path(filename)
    extension = file_path.suffix.lower()
    mime_type, _ = mimetypes.guess_type(filename)
    
    try:
        if extension == '.pdf':
            return _extract_pdf_text(file_content)
        elif extension == '.docx':
            return _extract_docx_text(file_content)
        elif extension == '.doc':
            raise DocumentParsingError("Legacy .doc format not supported. Please convert to .docx or PDF")
        elif extension == '.txt':
            return _extract_text_file(file_content)
        elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
            return _extract_image_text_ocr(file_content)
        elif mime_type and mime_type.startswith('image/'):
            return _extract_image_text_ocr(file_content)
        elif mime_type and mime_type.startswith('text/'):
            return _extract_text_file(file_content)
        else:
            raise DocumentParsingError(f"Unsupported file format: {extension} (MIME: {mime_type})")
    
    except DocumentParsingError:
        raise
    except Exception as e:
        logger.error(f"Unexpected error parsing file {filename}: {e}")
        raise DocumentParsingError(f"Failed to parse file: {str(e)}")

def _extract_pdf_text(file_content: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        
        if not text_parts:
            raise DocumentParsingError("PDF contains no extractable text")
        
        return "\n\n".join(text_parts)
    
    except fitz.FileDataError as e:
        raise DocumentParsingError(f"Invalid PDF file: {e}")
    except Exception as e:
        raise DocumentParsingError(f"PDF parsing failed: {e}")

def _extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            raise DocumentParsingError("DOCX contains no extractable text")
        
        return "\n".join(text_parts)
    
    except Exception as e:
        raise DocumentParsingError(f"DOCX parsing failed: {e}")

def _extract_text_file(file_content: bytes) -> str:
    """Extract text from plain text files."""
    try:
        # Try UTF-8 first
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to other encodings
            for encoding in ['latin1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # Last resort: decode with errors='replace'
            return file_content.decode('utf-8', errors='replace')
    
    except Exception as e:
        raise DocumentParsingError(f"Text file parsing failed: {e}")

def _extract_image_text_ocr(file_content: bytes) -> str:
    """Extract text from images using OCR (Tesseract)."""
    try:
        image = Image.open(io.BytesIO(file_content))
        
        # Convert to RGB if necessary (for PNG with transparency, etc.)
        if image.mode not in ['RGB', 'L']:
            image = image.convert('RGB')
        
        text = pytesseract.image_to_string(image)
        
        if not text.strip():
            raise DocumentParsingError("No text found in image")
        
        return text.strip()
    
    except pytesseract.TesseractNotFoundError:
        raise DocumentParsingError(
            "Tesseract OCR not installed. Please install tesseract-ocr system package."
        )
    except Exception as e:
        raise DocumentParsingError(f"OCR processing failed: {e}")

def get_supported_extensions() -> list[str]:
    """Get list of supported file extensions."""
    return ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']

def is_supported_file(filename: str) -> bool:
    """Check if file format is supported."""
    extension = Path(filename).suffix.lower()
    return extension in get_supported_extensions()