import io
import pytest
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

from orchestrator.doc_processing import (
    extract_text_from_file,
    DocumentParsingError,
    get_supported_extensions,
    is_supported_file,
    _extract_text_file,
    _extract_pdf_text,
    _extract_docx_text,
    _extract_image_text_ocr,
)


class TestDocumentProcessing:
    
    def test_get_supported_extensions(self):
        """Test getting list of supported extensions."""
        extensions = get_supported_extensions()
        expected = ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']
        assert set(extensions) == set(expected)
    
    def test_is_supported_file(self):
        """Test file support detection."""
        assert is_supported_file("test.pdf") is True
        assert is_supported_file("test.docx") is True
        assert is_supported_file("test.txt") is True
        assert is_supported_file("test.png") is True
        assert is_supported_file("test.jpg") is True
        assert is_supported_file("test.JPEG") is True  # Case insensitive
        
        assert is_supported_file("test.doc") is False
        assert is_supported_file("test.pptx") is False
        assert is_supported_file("test.exe") is False
    
    def test_extract_text_from_empty_file(self):
        """Test handling of empty files."""
        with pytest.raises(DocumentParsingError, match="Empty file content"):
            extract_text_from_file(b"", "test.txt")
    
    def test_extract_text_from_unsupported_format(self):
        """Test handling of unsupported file formats."""
        with pytest.raises(DocumentParsingError, match="Unsupported file format"):
            extract_text_from_file(b"some content", "test.xyz")
    
    def test_extract_text_file_utf8(self):
        """Test text file extraction with UTF-8 encoding."""
        content = "Hello, world! 🌍"
        content_bytes = content.encode('utf-8')
        
        result = extract_text_from_file(content_bytes, "test.txt")
        assert result == content
    
    def test_extract_text_file_latin1(self):
        """Test text file extraction with Latin-1 encoding."""
        content = "Héllo, wörld!"
        content_bytes = content.encode('latin1')
        
        result = extract_text_from_file(content_bytes, "test.txt")
        assert result == content
    
    def test_extract_text_file_direct(self):
        """Test direct text file extraction function."""
        content = "Sample text content"
        content_bytes = content.encode('utf-8')
        
        result = _extract_text_file(content_bytes)
        assert result == content
    
    def test_extract_text_file_with_encoding_fallback(self):
        """Test text extraction with encoding fallback."""
        # Create content with special characters that work in latin1 but not utf-8
        content_bytes = b'\xe9\xe8\xe0'  # éèà in latin1
        
        result = _extract_text_file(content_bytes)
        assert len(result) > 0  # Should decode something
    
    def test_legacy_doc_format_rejection(self):
        """Test that legacy .doc format is properly rejected."""
        with pytest.raises(DocumentParsingError, match="Legacy .doc format not supported"):
            extract_text_from_file(b"fake doc content", "test.doc")


class TestPDFProcessing:
    
    def test_pdf_text_extraction_simple(self):
        """Test simple PDF text extraction."""
        # Create a simple PDF with text
        import fitz
        
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 72), "Hello, PDF World!")
        
        pdf_bytes = doc.write()
        doc.close()
        
        result = extract_text_from_file(pdf_bytes, "test.pdf")
        assert "Hello, PDF World!" in result
    
    def test_pdf_multipage_extraction(self):
        """Test multi-page PDF text extraction."""
        import fitz
        
        doc = fitz.open()
        
        # Page 1
        page1 = doc.new_page()
        page1.insert_text((50, 72), "Page 1 content")
        
        # Page 2  
        page2 = doc.new_page()
        page2.insert_text((50, 72), "Page 2 content")
        
        pdf_bytes = doc.write()
        doc.close()
        
        result = extract_text_from_file(pdf_bytes, "test.pdf")
        assert "Page 1 content" in result
        assert "Page 2 content" in result
    
    def test_pdf_invalid_file(self):
        """Test handling of invalid PDF files."""
        with pytest.raises(DocumentParsingError, match="Invalid PDF file"):
            _extract_pdf_text(b"not a pdf file")
    
    def test_pdf_empty_content(self):
        """Test handling of PDF with no extractable text."""
        import fitz
        
        doc = fitz.open()
        doc.new_page()  # Empty page
        pdf_bytes = doc.write()
        doc.close()
        
        # Empty pages might still return whitespace characters, so we expect an exception
        try:
            result = _extract_pdf_text(pdf_bytes)
            # If no exception, verify the result is effectively empty
            assert result.strip() == ""
        except DocumentParsingError as e:
            # This is the expected behavior for truly empty PDFs
            assert "PDF contains no extractable text" in str(e)


class TestDOCXProcessing:
    
    def test_docx_text_extraction(self):
        """Test DOCX text extraction."""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("Hello, DOCX World!")
        doc.add_paragraph("Second paragraph")
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes = docx_bytes.getvalue()
        
        result = extract_text_from_file(docx_bytes, "test.docx")
        assert "Hello, DOCX World!" in result
        assert "Second paragraph" in result
    
    def test_docx_with_tables(self):
        """Test DOCX text extraction including tables."""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("Document with table:")
        
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes = docx_bytes.getvalue()
        
        result = extract_text_from_file(docx_bytes, "test.docx")
        assert "Document with table:" in result
        assert "Cell 1" in result
        assert "Cell 4" in result
    
    def test_docx_invalid_file(self):
        """Test handling of invalid DOCX files."""
        with pytest.raises(DocumentParsingError, match="DOCX parsing failed"):
            _extract_docx_text(b"not a docx file")
    
    def test_docx_empty_content(self):
        """Test handling of DOCX with no extractable text."""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("")  # Empty paragraph
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes = docx_bytes.getvalue()
        
        with pytest.raises(DocumentParsingError, match="DOCX contains no extractable text"):
            _extract_docx_text(docx_bytes)


def _tesseract_available() -> bool:
    """Check if Tesseract is available for testing."""
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except:
        return False


class TestImageOCRProcessing:
    
    def create_text_image(self, text: str, size=(200, 100)) -> bytes:
        """Helper to create an image with text for testing."""
        image = Image.new('RGB', size, color='white')
        draw = ImageDraw.Draw(image)
        
        # Use default font
        try:
            font = ImageFont.load_default()
        except:
            font = None
        
        # Calculate position to center text
        if font:
            bbox = draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
        else:
            text_width = len(text) * 10  # Rough estimate
            text_height = 15
        
        x = (size[0] - text_width) // 2
        y = (size[1] - text_height) // 2
        
        draw.text((x, y), text, fill='black', font=font)
        
        img_bytes = io.BytesIO()
        image.save(img_bytes, format='PNG')
        return img_bytes.getvalue()
    
    @pytest.mark.skipif(
        not _tesseract_available(),
        reason="Tesseract OCR not available"
    )
    def test_ocr_text_extraction(self):
        """Test OCR text extraction from images."""
        text = "Hello OCR"
        image_bytes = self.create_text_image(text)
        
        result = extract_text_from_file(image_bytes, "test.png")
        # OCR might not be perfect, but should contain the main text
        assert "Hello" in result or "OCR" in result
    
    @pytest.mark.skipif(
        not _tesseract_available(),
        reason="Tesseract OCR not available" 
    )
    def test_ocr_different_formats(self):
        """Test OCR with different image formats."""
        text = "Test"
        
        # Test PNG
        image_bytes = self.create_text_image(text)
        result_png = extract_text_from_file(image_bytes, "test.png")
        
        # Convert to JPEG
        image = Image.open(io.BytesIO(image_bytes))
        jpg_bytes = io.BytesIO()
        image.save(jpg_bytes, format='JPEG')
        
        result_jpg = extract_text_from_file(jpg_bytes.getvalue(), "test.jpg")
        
        # Both should extract some text
        assert len(result_png.strip()) > 0
        assert len(result_jpg.strip()) > 0
    
    def test_ocr_invalid_image(self):
        """Test handling of invalid image files."""
        with pytest.raises(DocumentParsingError):
            _extract_image_text_ocr(b"not an image")
    
    @pytest.mark.skipif(
        _tesseract_available(),
        reason="Tesseract is available - testing unavailable case"
    )
    def test_ocr_tesseract_not_installed(self):
        """Test handling when Tesseract is not installed."""
        image_bytes = self.create_text_image("Test")
        
        with pytest.raises(DocumentParsingError, match="Tesseract OCR not installed"):
            _extract_image_text_ocr(image_bytes)